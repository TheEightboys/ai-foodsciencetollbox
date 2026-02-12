"""
DOCX Export for Learning Objectives - IMPROVED VERSION
Exports to .docx with EXACT template preservation (fonts, spacing, styles, numbering).
"""

import os
import shutil
from docx import Document
from docx.shared import Pt


def export_to_docx_with_template_preservation(
    grade_level: str,
    topic: str,
    objectives: list,
    output_path: str,
    template_path: str
) -> str:
    """
    Export learning objectives to DOCX using template with EXACT formatting preservation.
    
    This method loads the template and modifies ONLY the content text while preserving
    all formatting, fonts, spacing, styles, and numbering from the template.
    
    Args:
        grade_level: Grade level string (e.g., "Middle", "High", "College")
        topic: Topic string (e.g., "Foodborne diseases")
        objectives: List of objective strings
        output_path: Path to save DOCX file
        template_path: Path to .docx template (REQUIRED)
    
    Returns:
        Path to saved file
    """
    
    if not template_path or not os.path.exists(template_path):
        raise ValueError(f"Template file required and must exist: {template_path}")
    
    # Load the template document
    doc = Document(template_path)
    
    # The ACTUAL template structure (from Template_-_Lesson_Objectives.docx):
    # Paragraph 0: "Lesson Objectives" (Title style)
    # Paragraph 1: "" (blank line, Normal style)
    # Paragraph 2: "Grade Level: Middle" (Normal style)
    # Paragraph 3: "Topic: Bacteria" (Normal style)
    # Paragraph 4: "" (blank line, Normal style)
    # Paragraph 5: "By the end of this lesson, students will be able to:" (Normal style)
    # Paragraph 6+: Numbered list items (List Number style)
    
    paragraphs = doc.paragraphs
    
    # Paragraph 0: "Lesson Objectives" - leave as is (Title style is fine)
    
    # Paragraph 1: blank line - leave as is
    
    # Paragraph 2: Update Grade Level
    if len(paragraphs) > 2:
        p = paragraphs[2]
        # Clear existing runs
        for run in p.runs:
            run.text = ''
        # Remove extra runs if any
        while len(p.runs) > 2:
            p._element.remove(p.runs[-1]._element)
        
        # Ensure we have at least 2 runs
        while len(p.runs) < 2:
            p.add_run()
        
        # Set the text with proper formatting
        p.runs[0].text = 'Grade Level: '
        p.runs[0].bold = True
        p.runs[1].text = grade_level
        p.runs[1].bold = False
    
    # Paragraph 3: Update Topic
    if len(paragraphs) > 3:
        p = paragraphs[3]
        # Clear existing runs
        for run in p.runs:
            run.text = ''
        # Remove extra runs if any
        while len(p.runs) > 2:
            p._element.remove(p.runs[-1]._element)
        
        # Ensure we have at least 2 runs
        while len(p.runs) < 2:
            p.add_run()
        
        # Set the text with proper formatting
        p.runs[0].text = 'Topic: '
        p.runs[0].bold = True
        p.runs[1].text = topic
        p.runs[1].bold = False
    
    # Paragraph 4: blank line - leave as is
    
    # Paragraph 5: "By the end..." - leave as is (already correct in template)
    
    # Paragraph 6+: Update objectives (List Number style)
    objective_start_index = 6
    
    for i, objective in enumerate(objectives):
        paragraph_index = objective_start_index + i
        
        if paragraph_index < len(paragraphs):
            # Update existing paragraph - preserve ALL formatting
            p = paragraphs[paragraph_index]
            # Clear text in all runs
            for run in p.runs:
                run.text = ''
            # Keep only the first run
            while len(p.runs) > 1:
                p._element.remove(p.runs[-1]._element)
            
            # Set the new objective text
            if len(p.runs) == 0:
                p.add_run()
            p.runs[0].text = objective
        else:
            # Need to add a new paragraph - copy style and formatting from first objective
            if objective_start_index < len(paragraphs):
                # Get the reference paragraph (first objective)
                reference_p = paragraphs[objective_start_index]
                
                # Add new paragraph with same style
                new_p = doc.add_paragraph(objective, style=reference_p.style)
                
                # Copy the paragraph properties (including numbering) from reference
                if reference_p._p.pPr is not None:
                    if new_p._p.pPr is None:
                        new_p._p.get_or_add_pPr()
                    
                    # Copy numbering properties
                    if reference_p._p.pPr.numPr is not None:
                        from copy import deepcopy
                        new_p._p.pPr._insert_numPr(deepcopy(reference_p._p.pPr.numPr))
                    
                    # Copy spacing properties
                    if reference_p._p.pPr.spacing is not None:
                        from copy import deepcopy
                        new_p._p.pPr._insert_spacing(deepcopy(reference_p._p.pPr.spacing))
    
    # Remove any extra objective paragraphs beyond what we need
    paragraphs_to_remove = []
    for i in range(len(objectives), 10):  # Check up to 10 objective slots
        paragraph_index = objective_start_index + i
        if paragraph_index < len(paragraphs):
            paragraphs_to_remove.append(paragraphs[paragraph_index])
        else:
            break
    
    # Remove extra paragraphs from the document
    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)
    
    # Save document
    doc.save(output_path)
    return output_path


def export_to_docx(
    grade_level: str,
    topic: str,
    objectives: list,
    output_path: str,
    template_path: str = None,
    title_style: str = 'Title',
    body_style: str = 'Normal',
    list_style: str = 'List Number'
) -> str:
    """
    Export learning objectives to DOCX file.
    
    RECOMMENDED: Always provide template_path for consistent formatting.
    
    Args:
        grade_level: Grade level string
        topic: Topic string
        objectives: List of objective strings
        output_path: Path to save DOCX file
        template_path: Path to .docx template (HIGHLY RECOMMENDED)
        title_style: Style name for title (used only if no template)
        body_style: Style name for body text (used only if no template)
        list_style: Style name for numbered list (used only if no template)
    
    Returns:
        Path to saved file
    """
    
    # If template is provided, use the template preservation method
    if template_path and os.path.exists(template_path):
        return export_to_docx_with_template_preservation(
            grade_level=grade_level,
            topic=topic,
            objectives=objectives,
            output_path=output_path,
            template_path=template_path
        )
    
    # Fallback: Create document from scratch (original method)
    doc = Document()
    
    # Add title
    p = doc.add_paragraph('Lesson Objectives')
    try:
        p.style = title_style
    except KeyError:
        # Fallback: make it look like a title
        run = p.runs[0]
        run.font.size = Pt(26)
        run.font.bold = True
    
    # Add blank line
    p = doc.add_paragraph()
    try:
        p.style = body_style
    except KeyError:
        pass
    
    # Add Grade Level (label bold)
    p = doc.add_paragraph()
    try:
        p.style = body_style
    except KeyError:
        pass
    run1 = p.add_run('Grade Level: ')
    run1.bold = True
    run2 = p.add_run(grade_level)
    run2.bold = False
    
    # Add Topic (label bold)
    p = doc.add_paragraph()
    try:
        p.style = body_style
    except KeyError:
        pass
    run1 = p.add_run('Topic: ')
    run1.bold = True
    run2 = p.add_run(topic)
    run2.bold = False
    
    # Add blank line
    p = doc.add_paragraph()
    try:
        p.style = body_style
    except KeyError:
        pass
    
    # Add "By the end..." line (whole line bold)
    p = doc.add_paragraph('By the end of this lesson, students will be able to:')
    try:
        p.style = body_style
    except KeyError:
        pass
    if p.runs:
        p.runs[0].bold = True
    
    # Add objectives as numbered list
    for objective in objectives:
        p = doc.add_paragraph(objective)
        try:
            p.style = list_style
        except KeyError:
            # Fallback: use default paragraph style
            pass
    
    # Save document
    doc.save(output_path)
    return output_path


def export_learning_objectives_to_docx(
    grade_level: str,
    topic: str,
    objectives: list,
    output_path: str,
    template_path: str = None,
    title_style: str = 'Title',
    body_style: str = 'Normal',
    list_style: str = 'List Number'
) -> str:
    """
    Convenience wrapper for exporting learning objectives to DOCX.
    
    IMPORTANT: For consistent formatting, always provide template_path.
    
    Args:
        grade_level: Grade level string
        topic: Topic string
        objectives: List of objective strings
        output_path: Path to save DOCX file
        template_path: Path to .docx template (RECOMMENDED)
        title_style: Style name for title (used only without template)
        body_style: Style name for body text (used only without template)
        list_style: Style name for numbered list (used only without template)
    
    Returns:
        Path to saved file
    """
    return export_to_docx(
        grade_level,
        topic,
        objectives,
        output_path,
        template_path,
        title_style,
        body_style,
        list_style
    )
