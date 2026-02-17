"""
Document formatting service for generating print-ready DOCX and PDF files.

This service ensures all generated documents follow the exact formatting specifications:
- Arial 20pt bold centered for titles
- Arial 12pt for body text
- Arial 14pt bold for section headers
- Proper spacing (8pt after title, 6pt before section headings, 4pt after section heading, 6pt after heading)
- DOCX and PDF export support

Libraries used:
- python-docx: Industry standard for DOCX generation (best choice)
- fpdf2: PDF generation - pure Python, no system dependencies (works in containers)
"""

import io
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

# PDF generation using FPDF2 (pure Python, no system dependencies - works in containers)
try:
    from fpdf import FPDF
    FPDF2_AVAILABLE = True
except ImportError:
    FPDF2_AVAILABLE = False
    FPDF = None


class DocumentFormatter:
    """
    Formats generated content into print-ready DOCX and PDF documents
    following the exact specifications from the universal templates.
    """
    
    # Font specifications
    TITLE_FONT = 'Arial'
    TITLE_SIZE = Pt(20)
    BODY_FONT = 'Arial'
    BODY_SIZE = Pt(12)
    HEADER_FONT = 'Arial'
    HEADER_SIZE = Pt(14)
    
    # Spacing specifications (in points)
    SPACE_AFTER_TITLE = Pt(8)
    SPACE_BEFORE_SECTION = Pt(6)
    SPACE_AFTER_SECTION_HEADER = Pt(4)
    SPACE_AFTER_HEADING = Pt(6)
    
    def __init__(self):
        self.docx_doc = None
        self.pdf_buffer = None
    
    def _capitalize_grade_level(self, grade_level: str) -> str:
        """Capitalize grade level properly (e.g., 'middle' -> 'Middle')."""
        if not grade_level:
            return grade_level
        # Handle common grade level formats
        grade_lower = grade_level.lower().strip()
        grade_mapping = {
            'elementary': 'Elementary',
            'middle': 'Middle',
            'middle school': 'Middle School',
            'high': 'High',
            'high school': 'High School',
            'college': 'College',
        }
        # Check for exact matches first
        if grade_lower in grade_mapping:
            return grade_mapping[grade_lower]
        # Otherwise capitalize first letter
        return grade_level.capitalize()
    
    def _capitalize_topic(self, topic: str) -> str:
        """Capitalize topic properly (e.g., 'protein' -> 'Protein')."""
        if not topic:
            return topic
        # Capitalize first letter only
        return topic[0].upper() + topic[1:] if len(topic) > 1 else topic.upper()
    
    def _strip_metadata_from_content(self, content: str) -> str:
        """Remove metadata lines (Grade Level, Time Needed, Topic, Category) and generation type titles from AI-generated content."""
        import re
        lines = content.split('\n')
        filtered_lines = []
        skip_next_empty = False  # Flag to skip empty line after metadata
        
        for i, line in enumerate(lines):
            line_stripped = line.strip()
            line_lower = line_stripped.lower()
            
            # Skip lines that look like metadata (case-insensitive, with or without value)
            # Match patterns like:
            # - "Grade Level: middle"
            # - "**Grade Level**: middle" (new format)
            # - "Grade Level:"
            # - "Time Needed: 3-5 minutes"
            # - "Topic: protein"
            # - "**Topic**: protein" (new format)
            # - "Category: food_science"
            if (re.match(r'^(\*\*)?(grade level|time needed|topic|category)(\*\*)?\s*:\s*', line_lower) or
                re.match(r'^(\*\*)?(grade level|time needed|topic|category)(\*\*)?\s*$', line_lower)):
                skip_next_empty = True
                continue
            
            # Also check for lines that start with just the label
            if line_lower in ['grade level', 'time needed', 'topic', 'category']:
                skip_next_empty = True
                continue
            
            # Skip "Purpose" lines that appear in discussion questions
            if line_lower == 'purpose':
                skip_next_empty = True
                continue
            
            # Skip generation type titles that appear after metadata
            # "Lesson Objectives", "Lesson Starter", "Lesson Starter Ideas",
            # "Bell Ringer", "Today's Bell Ringer", "Discussion Questions"
            if (line_lower == 'lesson objectives' or 
                line_lower == 'lesson starter' or 
                line_lower == 'lesson starter ideas' or
                line_lower == 'discussion questions' or
                line_lower == "bell ringer" or
                line_lower == "today's bell ringer" or
                line_lower.startswith('lesson objectives') or
                line_lower.startswith('lesson starter ideas') or
                line_lower.startswith('lesson starter') or
                line_lower.startswith('discussion questions') or
                line_lower.startswith("bell ringer") or
                line_lower.startswith("today's bell ringer")):
                # Only skip if it's a short line (likely metadata, not part of content)
                if len(line_stripped) < 50:
                    skip_next_empty = True
                    continue
            
            # Remove "(Section header: ...)" text that AI might include
            # This appears in prompts like "(Section header: Arial 14pt, bold)"
            if re.search(r'\(section header[^)]*\)', line_lower):
                # Remove the entire line if it's just the section header instruction
                if re.match(r'^\s*\(section header[^)]*\)\s*$', line_lower):
                    continue
                # Otherwise, remove just the "(Section header: ...)" part from the line
                line = re.sub(r'\(section header[^)]*\)', '', line, flags=re.IGNORECASE).strip()
                line_stripped = line.strip()
                line_lower = line_stripped.lower()
            
            # Skip empty lines if we just removed metadata
            if skip_next_empty and line_stripped == '':
                skip_next_empty = False
                continue
            
            # Reset skip flag if we encounter non-empty line
            if line_stripped != '':
                skip_next_empty = False
            
            # Keep the line
            filtered_lines.append(line)
        
        return '\n'.join(filtered_lines).strip()
    
    def _add_metadata_paragraph(self, doc: Document, label: str, value: str):
        """Add a metadata paragraph with bold label and capitalized value."""
        para = doc.add_paragraph()
        # Add bold label
        label_run = para.add_run(f'{label}: ')
        label_run.bold = True
        label_run.font.name = self.BODY_FONT
        label_run.font.size = self.BODY_SIZE
        # Add value
        value_run = para.add_run(value)
        value_run.font.name = self.BODY_FONT
        value_run.font.size = self.BODY_SIZE
    
    def format_lesson_starter(self, topic: str, grade_level: str, content: str, subject: str = "Food Science") -> Dict[str, Any]:
        """
        Format lesson starter ideas (v2) into a clean, printable DOCX
        matching the exact template screenshot.

        Structure:
        0. Title: "Lesson Starter Ideas" – centered, bold, underlined, Arial 20pt
        1. Blank line
        2. Grade Level: {grade_level}  – bold dark-teal label, regular value
        3. Topic: {topic}              – bold dark-teal label, regular value
        4. Blank line
        5-11. Seven options, each:
              - "Option N:" bold dark-teal + title bold black
              - 3-5 sentence description as regular body text
        """
        import re
        doc = Document()

        DARK_TEAL = RGBColor(0x1F, 0x49, 0x7D)

        # ── Default Normal style ────────────────────────────────
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.space_after = Pt(0)
        normal_style.paragraph_format.space_before = Pt(0)
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_style.paragraph_format.line_spacing = 1.15

        def _set_spacing(para):
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.15

        def _blank():
            bp = doc.add_paragraph()
            _set_spacing(bp)

        # ── 0  Title – centered, bold, underlined, 20pt ─────────
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(title_para)
        title_run = title_para.add_run('Lesson Starter Ideas')
        title_run.bold = True
        title_run.underline = True
        title_run.font.size = Pt(20)
        title_run.font.name = 'Arial'

        # ── 1  Blank ────────────────────────────────────────────
        _blank()

        # ── 2  Grade Level (bold dark-teal label) ───────────────
        capitalized_grade = self._capitalize_grade_level(grade_level)
        gl_para = doc.add_paragraph()
        _set_spacing(gl_para)
        gl_label = gl_para.add_run('Grade Level: ')
        gl_label.bold = True
        gl_label.font.name = 'Arial'
        gl_label.font.size = Pt(12)
        gl_label.font.color.rgb = DARK_TEAL
        gl_value = gl_para.add_run(capitalized_grade)
        gl_value.font.name = 'Arial'
        gl_value.font.size = Pt(12)

        # ── 3  Topic (bold dark-teal label) ─────────────────────
        capitalized_topic = self._capitalize_topic(topic)
        tp_para = doc.add_paragraph()
        _set_spacing(tp_para)
        tp_label = tp_para.add_run('Topic: ')
        tp_label.bold = True
        tp_label.font.name = 'Arial'
        tp_label.font.size = Pt(12)
        tp_label.font.color.rgb = DARK_TEAL
        tp_value = tp_para.add_run(capitalized_topic)
        tp_value.font.name = 'Arial'
        tp_value.font.size = Pt(12)

        # ── 4  Blank ────────────────────────────────────────────
        _blank()

        # ── 5-11  Parse and render 7 ideas ──────────────────────
        cleaned_content = self._strip_metadata_from_content(content)

        # Extract ideas — support both "Option N:" and legacy "Starter Idea N:"
        idea_pattern = re.compile(
            r'(?:Option|Starter\s+Idea)\s+(\d+)\s*:\s*(.+?)(?=\n(?:Option|Starter\s+Idea)\s+\d+\s*:|\Z)',
            re.DOTALL | re.IGNORECASE,
        )
        ideas = list(idea_pattern.finditer(cleaned_content))

        if ideas:
            for match in ideas:
                number = match.group(1)
                block = match.group(2).strip()
                # First line is the title, rest is description
                lines = block.split('\n', 1)
                title_text = lines[0].strip()
                description = lines[1].strip() if len(lines) > 1 else ""

                # ── Blank line before each option (visual spacing) ──
                _blank()

                # Option heading: "Option N:" bold dark-teal + title bold black
                idea_para = doc.add_paragraph()
                _set_spacing(idea_para)
                label_run = idea_para.add_run(f'Option {number}:  ')
                label_run.bold = True
                label_run.font.name = 'Arial'
                label_run.font.size = Pt(12)
                label_run.font.color.rgb = DARK_TEAL
                title_run2 = idea_para.add_run(title_text)
                title_run2.bold = True
                title_run2.font.name = 'Arial'
                title_run2.font.size = Pt(12)

                # Description: regular body text
                if description:
                    desc_para = doc.add_paragraph()
                    _set_spacing(desc_para)
                    desc_run = desc_para.add_run(description)
                    desc_run.font.name = 'Arial'
                    desc_run.font.size = Pt(12)
        else:
            # Fallback: if regex didn't match, dump content as-is
            for line in cleaned_content.split('\n'):
                stripped = line.strip()
                if not stripped:
                    _blank()
                    continue
                p = doc.add_paragraph()
                _set_spacing(p)
                r = p.add_run(stripped)
                r.font.name = 'Arial'
                r.font.size = Pt(12)

        return self._create_docx_buffer(doc)
    
    def format_learning_objectives(self, topic: str, grade_level: str, content: str, subject: str = "Food Science") -> Dict[str, Any]:
        """
        Format learning objectives to match the exact template screenshot.

          Structure:
          0. Title: "Learning Objectives" – left-aligned, underlined, dark teal,
              Arial 14pt bold
          1. Blank line
          2. Topic: {topic} – Arial 14pt, label bold
          3. Grade: {grade_level} – Arial 14pt, label bold
          4. Blank line
          5. "By the end of this lesson, students will be able to:" – bold, Arial 12pt
          6+. Numbered objectives – decimal "1.", indented, Arial 12pt, no bold
        """
        import re
        doc = Document()

        # ── default Normal style ──────────────────────────────────
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.space_after = Pt(0)
        normal_style.paragraph_format.space_before = Pt(0)
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_style.paragraph_format.line_spacing = 1.5

        # ── helper: apply 1.5 line spacing to a paragraph ─────────
        def _set_spacing(para):
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.5

        # ── helper: blank paragraph ───────────────────────────────
        def _blank():
            bp = doc.add_paragraph()
            _set_spacing(bp)

        # ── 0  Title – left-aligned, underlined, dark teal ───────
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_spacing(title_para)
        title_run = title_para.add_run('Learning Objectives')
        title_run.bold = True
        title_run.underline = True
        title_run.font.size = Pt(14)
        title_run.font.name = 'Arial'
        title_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)  # dark teal

        # ── 1  Blank line ────────────────────────────────────────
        _blank()

        # ── 2  Topic ─────────────────────────────────────────────
        capitalized_topic = self._capitalize_topic(topic)
        topic_para = doc.add_paragraph()
        _set_spacing(topic_para)
        t_label = topic_para.add_run('Topic: ')
        t_label.bold = True
        t_label.font.name = 'Arial'
        t_label.font.size = Pt(14)
        t_value = topic_para.add_run(capitalized_topic)
        t_value.bold = False
        t_value.font.name = 'Arial'
        t_value.font.size = Pt(14)

        # ── 3  Grade ───────────────────────────────────────
        capitalized_grade = self._capitalize_grade_level(grade_level)
        grade_para = doc.add_paragraph()
        _set_spacing(grade_para)
        g_label = grade_para.add_run('Grade: ')
        g_label.bold = True
        g_label.font.name = 'Arial'
        g_label.font.size = Pt(14)
        g_value = grade_para.add_run(capitalized_grade)
        g_value.bold = False
        g_value.font.name = 'Arial'
        g_value.font.size = Pt(14)

        # ── 4  Blank line ────────────────────────────────────────
        _blank()

        # ── 5  Stem sentence (entire line bold) ──────────────────
        stem_para = doc.add_paragraph()
        _set_spacing(stem_para)
        stem_run = stem_para.add_run('By the end of this lesson, students will be able to:')
        stem_run.bold = True
        stem_run.font.name = 'Arial'
        stem_run.font.size = Pt(12)

        # ── Extract objectives from AI content ───────────────────
        cleaned_content = self._strip_metadata_from_content(content)
        sections = self._parse_learning_objectives_sections(cleaned_content)
        objectives_content = sections.get('Learning Objectives', cleaned_content)

        lines = objectives_content.split('\n')
        objectives = []
        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            if 'by the end' in line_stripped.lower() or 'students will be able to' in line_stripped.lower():
                continue
            if line_stripped.lower().startswith('category:'):
                continue
            if (':' in line_stripped and
                    line_stripped.endswith(':') and
                    any(kw in line_stripped.lower() for kw in
                        ['analyzing', 'applying', 'using', 'demonstrating', 'students will'])):
                continue
            clean_line = re.sub(r'^\d+[\.\)\s]+', '', line_stripped)
            clean_line = re.sub(r'^[•\-\*]\s+', '', clean_line)
            if clean_line.lower().startswith('category') or clean_line.lower().startswith('topic'):
                continue
            if clean_line and len(clean_line.strip()) > 3 and not clean_line.strip().endswith(':'):
                objectives.append(clean_line.strip())

        # ── 3+  Numbered objectives ──────────────────────────────
        # Use 'List Number' style for proper decimal numbering,
        # then override indent to match template: left 0.5 in, hanging 0.25 in.
        # Content is Arial 12pt, NO bold.
        for obj in objectives:
            if not obj:
                continue
            p = doc.add_paragraph(style='List Number')
            _set_spacing(p)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.25)

            obj_run = p.add_run(obj)
            obj_run.font.name = 'Arial'
            obj_run.font.size = Pt(12)

        return self._create_docx_buffer(doc)
    
    
    def _parse_lesson_starter_sections(self, content: str) -> Dict[str, str]:
        """Parse lesson starter content into sections."""
        sections = {}
        current_section = None
        current_content = []
        
        lines = content.split('\n')
        
        section_headers = [
            'Key Lesson Ideas to Explore',
            'Prior Knowledge to Connect',
            'Teacher Opening Script (3–4 minutes)',
            # Legacy headers for backward compatibility
            'Hook',
            'Why This Lesson Matters',
            'Background Connection',
            'Key Ideas to Explore',
            'Introductory Teacher Script'
        ]
        
        for line in lines:
            original_line = line
            line = line.strip()
            
            # Check if this line is a section header (check BEFORE handling empty lines)
            is_header = False
            for header in section_headers:
                # More precise header matching
                if (line == header or 
                    (line.lower() == header.lower()) or
                    (line.lower().startswith(header.lower()) and len(line) < len(header) + 10 and not line[len(header):].strip().startswith(':'))):
                    is_header = True
                    # Save previous section if it exists
                    if current_section and current_content:
                        sections[current_section] = '\n'.join(current_content).strip()
                    # Start new section
                    current_section = header
                    current_content = []
                    break
            
            if is_header:
                continue  # Skip the header line itself
            
            # Handle empty lines - preserve them within sections but don't end the section
            if not line:
                if current_section:
                    # Preserve empty lines within the section (they might be formatting)
                    current_content.append('')
                continue
            
            # Add content to current section
            if current_section:
                current_content.append(original_line)
            elif not sections:
                # Content before first section - should be Key Lesson Ideas to Explore
                current_section = 'Key Lesson Ideas to Explore'
                current_content.append(original_line)
        
        # Add last section (important - don't lose the final section)
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        # Fallback: if no sections found, try to preserve content
        if not sections and content.strip():
            # Try to find any section headers in the content
            for header in section_headers:
                if header.lower() in content.lower():
                    # Found a header, treat everything as one section
                    sections['Hook'] = content.strip()
                    break
            # If still no sections, preserve all content as Hook
            if not sections:
                sections['Hook'] = content.strip()
        
        return sections
    
    def _parse_learning_objectives_sections(self, content: str) -> Dict[str, str]:
        """Parse learning objectives content into sections."""
        sections = {}
        current_section = None
        current_content = []
        
        lines = content.split('\n')
        
        section_headers = ['Learning Objectives']  # Success Criteria has been removed
        
        for line in lines:
            original_line = line
            line = line.strip()
            
            # Check if this line is a section header (check BEFORE handling empty lines)
            is_header = False
            for header in section_headers:
                # More precise header matching - must be exact match or start with header and be short
                if (line == header or 
                    (line.lower() == header.lower()) or
                    (line.lower().startswith(header.lower()) and len(line) < len(header) + 10 and not line[len(header):].strip().startswith(':'))):
                    is_header = True
                    # Save previous section if it exists
                    if current_section and current_content:
                        sections[current_section] = '\n'.join(current_content).strip()
                    # Start new section
                    current_section = header
                    current_content = []
                    break
            
            if is_header:
                continue  # Skip the header line itself
            
            # Handle empty lines - preserve them within sections but don't end the section
            if not line:
                if current_section:
                    # Preserve empty lines within the section (they might be formatting)
                    current_content.append('')
                continue
            
            # Add content to current section
            if current_section:
                current_content.append(original_line)
            elif not sections:
                # Content before first section - should be Learning Objectives
                current_section = 'Learning Objectives'
                current_content.append(original_line)
        
        # Add last section (important - don't lose the final section)
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        # Fallback: if no sections found, treat entire content as Learning Objectives
        if not sections and content.strip():
            sections['Learning Objectives'] = content.strip()
        
        return sections
    
    def _parse_bell_ringer_sections(self, content: str) -> Dict[str, str]:
        """Parse bell ringer content into sections."""
        import re
        sections = {}
        current_section = None
        current_content = []
        
        lines = content.split('\n')
        
        section_headers = ['Purpose', 'Prompt']
        
        for line in lines:
            original_line = line
            line = line.strip()
            
            # Remove "(Section header: ...)" text from lines
            if re.search(r'\(section header[^)]*\)', line, re.IGNORECASE):
                # Remove the entire line if it's just the section header instruction
                if re.match(r'^\s*\(section header[^)]*\)\s*$', line, re.IGNORECASE):
                    continue
                # Otherwise, remove just the "(Section header: ...)" part from the line
                line = re.sub(r'\(section header[^)]*\)', '', line, flags=re.IGNORECASE).strip()
                original_line = line  # Update original_line to cleaned version
            
            # Check if this line is a section header (check BEFORE handling empty lines)
            is_header = False
            for header in section_headers:
                # More precise header matching
                if (line == header or 
                    (line.lower() == header.lower()) or
                    (line.lower().startswith(header.lower()) and len(line) < len(header) + 10 and not line[len(header):].strip().startswith(':'))):
                    is_header = True
                    # Save previous section if it exists
                    if current_section and current_content:
                        sections[current_section] = '\n'.join(current_content).strip()
                    # Start new section
                    current_section = header
                    current_content = []
                    break
            
            if is_header:
                continue  # Skip the header line itself
            
            # Handle empty lines - preserve them within sections but don't end the section
            if not line:
                if current_section:
                    # Preserve empty lines within the section (they might be formatting)
                    current_content.append('')
                continue
            
            # If we have a current section, add the line (including empty lines)
            if current_section:
                current_content.append(original_line)
            elif not sections:
                # Content before first section - should be Purpose
                current_section = 'Purpose'
                current_content.append(original_line)
        
        # Add last section (important - don't lose the final section)
        if current_section and current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        # Fallback: if no sections found, try to preserve content
        if not sections and content.strip():
            # Try to find any section headers in the content
            for header in section_headers:
                if header.lower() in content.lower():
                    # Found a header, treat everything as one section
                    sections['Purpose'] = content.strip()
                    break
            # If still no sections, preserve all content as Purpose
            if not sections:
                sections['Purpose'] = content.strip()
        
        return sections
    

    
    def format_discussion_questions(self, topic: str, grade_level: str, content: str, subject: str = "Food Science") -> Dict[str, Any]:
        """
        Format discussion questions (v2) into a clean, printable DOCX
        matching the exact template.

        Structure:
        0. Title: "Discussion Questions" — centered, bold, italic, Arial 20pt
        1. Blank line
        2. Grade Level: {grade_level} — bold dark-teal label, regular value
        3. Topic: {topic} — bold dark-teal label, regular value  (NO blank between)
        4. Blank line
        5-9. Five options, each:
             - "Option N" bold dark-teal on its own line
             - Question paragraph as regular body text
             - "Teacher cue: ..." all regular weight (NOT bold)
        """
        import re
        doc = Document()

        DARK_TEAL = RGBColor(0x1F, 0x49, 0x7D)

        # ── Default Normal style ────────────────────────────
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.space_after = Pt(0)
        normal_style.paragraph_format.space_before = Pt(0)
        normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        normal_style.paragraph_format.line_spacing = 1.15

        def _set_spacing(para):
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            para.paragraph_format.line_spacing = 1.15

        def _blank():
            bp = doc.add_paragraph()
            _set_spacing(bp)

        # ── 0  Title — centered, bold, italic, 20pt ─────────
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(title_para)
        title_run = title_para.add_run('Discussion Questions')
        title_run.bold = True
        title_run.italic = True
        title_run.font.size = Pt(20)
        title_run.font.name = 'Arial'

        # ── 1  Blank ────────────────────────────────────────
        _blank()

        # ── 2  Grade Level ──────────────────────────────────
        capitalized_grade = self._capitalize_grade_level(grade_level)
        gl_para = doc.add_paragraph()
        _set_spacing(gl_para)
        gl_label = gl_para.add_run('Grade Level: ')
        gl_label.bold = True
        gl_label.font.name = 'Arial'
        gl_label.font.size = Pt(12)
        gl_label.font.color.rgb = DARK_TEAL
        gl_value = gl_para.add_run(capitalized_grade)
        gl_value.font.name = 'Arial'
        gl_value.font.size = Pt(12)

        # ── 3  Topic (immediately after Grade Level, no blank line) ──
        capitalized_topic = self._capitalize_topic(topic)
        tp_para = doc.add_paragraph()
        _set_spacing(tp_para)
        tp_label = tp_para.add_run('Topic: ')
        tp_label.bold = True
        tp_label.font.name = 'Arial'
        tp_label.font.size = Pt(12)
        tp_label.font.color.rgb = DARK_TEAL
        tp_value = tp_para.add_run(capitalized_topic)
        tp_value.font.name = 'Arial'
        tp_value.font.size = Pt(12)

        # ── 4  Blank ────────────────────────────────────────
        _blank()

        # ── 6-10  Parse and render 5 option blocks ──────────
        cleaned_content = self._strip_metadata_from_content(content)

        # Extract "Option N" blocks
        option_pattern = re.compile(
            r'Option\s+(\d+)\s*\n(.+?)(?=Option\s+\d+\s*\n|\Z)',
            re.DOTALL | re.IGNORECASE,
        )
        options = list(option_pattern.finditer(cleaned_content))

        if options:
            for match in options:
                number = match.group(1)
                body = match.group(2).strip()

                # Separate question from Teacher cue
                cue_match = re.search(
                    r'Teacher\s+cue:\s*(.+)', body, re.IGNORECASE,
                )
                if cue_match:
                    question_text = body[: cue_match.start()].strip()
                    cue_text = cue_match.group(1).strip()
                else:
                    question_text = body
                    cue_text = ""

                # Blank line before each option (visual spacing)
                _blank()

                # Option heading: "Option N" bold dark-teal on own line
                opt_para = doc.add_paragraph()
                _set_spacing(opt_para)
                opt_run = opt_para.add_run(f'Option {number}')
                opt_run.bold = True
                opt_run.font.name = 'Arial'
                opt_run.font.size = Pt(12)
                opt_run.font.color.rgb = DARK_TEAL

                # Question paragraph: regular body text
                if question_text:
                    q_para = doc.add_paragraph()
                    _set_spacing(q_para)
                    q_run = q_para.add_run(question_text)
                    q_run.font.name = 'Arial'
                    q_run.font.size = Pt(12)

                # Teacher cue: all regular weight (matches template)
                if cue_text:
                    cue_para = doc.add_paragraph()
                    _set_spacing(cue_para)
                    cue_run = cue_para.add_run(f'Teacher cue: {cue_text}')
                    cue_run.font.name = 'Arial'
                    cue_run.font.size = Pt(12)
        else:
            # Fallback: dump content as-is line by line
            for line in cleaned_content.split('\n'):
                stripped = line.strip()
                if not stripped:
                    _blank()
                    continue
                p = doc.add_paragraph()
                _set_spacing(p)
                r = p.add_run(stripped)
                r.font.name = 'Arial'
                r.font.size = Pt(12)

        return self._create_docx_buffer(doc)
    
    def _extract_objectives_from_content(self, content: str) -> list:
        """Extract objectives from content, handling numbered lists and multi-line objectives.
        Filters out non-objective closing sentences and conversational text."""
        import re
        objectives = []
        lines = content.split('\n')
            
        current_objective = None
        skip_next_empty = False
            
        # Patterns that indicate non-objective conversational/closing text
        non_objective_patterns = [
            r'^(these objectives|these goals|these learning objectives|remember|note that|keep in mind|as you|when you|you can|you will|this lesson|the lesson|today|we will|we\'ll|let\'s|it\'s important|in conclusion|to summarize|finally)',
            r'^(students should|teachers should|educators should)',
            r'^by (completing|finishing|the end of)',
        ]
            
        for line in lines:
            original_line = line
            line = line.strip()
                
            # Skip the "By the end of this lesson" line
            if 'by the end' in line.lower() or 'students will be able to' in line.lower():
                skip_next_empty = True
                continue
                
            # Skip empty line after "By the end" if needed
            if skip_next_empty and not line:
                skip_next_empty = False
                continue
            skip_next_empty = False
                
            if not line:
                # Empty line - end current objective if exists
                if current_objective:
                    objectives.append(current_objective.strip())
                    current_objective = None
                continue
                
            # Check if this line is non-objective conversational text
            line_lower = line.lower()
            is_non_objective = False
            for pattern in non_objective_patterns:
                if re.match(pattern, line_lower):
                    is_non_objective = True
                    break
                
            # Also check for conversational phrases that indicate closing text
            conversational_phrases = [
                'remember that', 'note that', 'keep in mind', 'as you', 'when you',
                'this lesson will', 'today we will', 'we\'ll', 'let\'s', 'it\'s important',
                'in conclusion', 'to summarize', 'finally', 'these objectives help',
                'these goals', 'by completing', 'by finishing'
            ]
            for phrase in conversational_phrases:
                if phrase in line_lower and not re.match(r'^\d+[\.]\)?\s*', line):
                    is_non_objective = True
                    break
                
            # If we detect non-objective text, stop extracting (save current objective first)
            if is_non_objective:
                if current_objective:
                    objectives.append(current_objective.strip())
                    current_objective = None
                # Stop processing - we've hit conversational/closing text
                break
                
            # Check if this line starts a new numbered objective
            # Patterns: "1. ", "1)", "1 ", or just starts with a number followed by punctuation
            numbered_match = re.match(r'^(\d+)[\.]\)\s*(.+)$', line)
            if numbered_match:
                # Save previous objective if exists
                if current_objective:
                    objectives.append(current_objective.strip())
                # Start new objective
                current_objective = numbered_match.group(2).strip()
            elif re.match(r'^\d+\s+', line):
                # Number followed by space (no punctuation)
                match = re.match(r'^(\d+)\s+(.+)$', line)
                if match:
                    if current_objective:
                        objectives.append(current_objective.strip())
                    current_objective = match.group(2).strip()
            elif current_objective:
                # Continuation of current objective (multi-line objective)
                # Only continue if it looks like part of the objective (not a new sentence starting with capital)
                # Check if line starts with capital letter and is a complete sentence (ends with period)
                if re.match(r'^[A-Z][^.!?]*[.!?]\s*$', line):
                    # This looks like a complete sentence - might be end of objective or start of new text
                    # If current objective doesn't end with punctuation, this is likely continuation
                    if not current_objective.rstrip().endswith(('.', '!', '?')):
                        current_objective += ' ' + line
                    else:
                        # Current objective is complete, this might be a new objective or closing text
                        objectives.append(current_objective.strip())
                        # Check if this looks like an objective (starts with action verb)
                        action_verbs = ['identify', 'list', 'describe', 'explain', 'demonstrate', 'analyze', 
                                       'compare', 'apply', 'evaluate', 'justify', 'assess', 'propose', 'synthesize',
                                       'name', 'show', 'tell', 'create', 'develop', 'design']
                        if any(line_lower.startswith(verb + ' ') for verb in action_verbs):
                            current_objective = line
                        else:
                            current_objective = None
                            break
                else:
                    # Continuation of current objective
                    current_objective += ' ' + line
            else:
                # No current objective and doesn't start with number
                # Only accept if it starts with an action verb (likely an objective)
                action_verbs = ['identify', 'list', 'describe', 'explain', 'demonstrate', 'analyze', 
                               'compare', 'apply', 'evaluate', 'justify', 'assess', 'propose', 'synthesize',
                               'name', 'show', 'tell', 'create', 'develop', 'design']
                line_lower_start = line_lower.split()[0] if line_lower.split() else ''
                if any(line_lower_start.startswith(verb) for verb in action_verbs):
                    clean_line = line.lstrip('•-*0123456789. )')
                    if clean_line and len(clean_line) > 5:
                        current_objective = clean_line
                else:
                    # Doesn't look like an objective - stop processing
                    break
            
        # Don't forget the last objective
        if current_objective:
            objectives.append(current_objective.strip())
            
        # Filter out any objectives that look like conversational text
        filtered_objectives = []
        for obj in objectives:
            obj_lower = obj.lower().strip()
            # Skip if it's clearly not an objective
            if not any(obj_lower.startswith(verb + ' ') for verb in ['identify', 'list', 'describe', 'explain', 
                                                                     'demonstrate', 'analyze', 'compare', 'apply', 
                                                                     'evaluate', 'justify', 'assess', 'propose', 
                                                                     'synthesize', 'name', 'show', 'tell', 'create', 
                                                                     'develop', 'design']):
                # Check if it's conversational
                if not any(phrase in obj_lower for phrase in ['remember', 'note', 'keep in mind', 'this lesson', 
                                                              'today', 'we will', 'let\'s', 'it\'s important']):
                    filtered_objectives.append(obj)
            else:
                filtered_objectives.append(obj)
            
        # Fallback: if no objectives found, try to extract any meaningful lines
        if not filtered_objectives and content.strip():
            for line in content.split('\n'):
                line = line.strip()
                if line and 'by the end' not in line.lower() and 'students will be able to' not in line.lower():
                    # Check if it starts with a number or action verb
                    if re.match(r'^\d+[\.]\)?\s*', line) or any(line.lower().startswith(verb + ' ') 
                                                               for verb in ['identify', 'list', 'describe', 'explain']):
                        clean_line = line.lstrip('•-*0123456789. )')
                        if clean_line and len(clean_line) > 5:
                            filtered_objectives.append(clean_line)
            
        return filtered_objectives if filtered_objectives else []
    
    def _extract_bullet_points(self, content: str) -> list:
        """Extract bullet points from content."""
        points = []
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Skip the "Students will demonstrate success" line
            if 'students will demonstrate success' in line.lower():
                continue
            
            # Extract bullet points
            if line.startswith('•') or line.startswith('-') or line.startswith('*'):
                points.append(line)
            elif len(line) > 10:
                # Might be a continuation or standalone point
                points.append(line)
        
        return points if points else [content]
    
    def _create_docx_buffer(self, doc: Document) -> Dict[str, Any]:
        """Create a buffer with the DOCX file."""
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return {'docx': buffer, 'format': 'docx'}
    
    @staticmethod
    def _sanitize_for_pdf(text: str) -> str:
        """Sanitize text for FPDF2 Latin-1 built-in fonts.
        
        FPDF2 built-in fonts (Helvetica, Courier, Times) only support Latin-1.
        This replaces common Unicode characters with their ASCII equivalents
        so the PDF renders correctly without crashing.
        """
        if not text:
            return text
        replacements = {
            # Bullets & dashes
            '\u2022': '-',   # •
            '\u2023': '>',   # ‣
            '\u2043': '-',   # ⁃
            '\u2013': '-',   # –
            '\u2014': '-',   # —
            '\u2015': '-',   # ―
            # Smart quotes
            '\u2018': "'",  # '
            '\u2019': "'",  # '
            '\u201A': "'",  # ‚
            '\u201C': '"',  # "
            '\u201D': '"',  # "
            '\u201E': '"',  # „
            '\u201F': '"',  # ‟
            '\u2039': '<',   # ‹
            '\u203A': '>',   # ›
            '\u00AB': '"',  # «
            '\u00BB': '"',  # »
            # Spaces & special
            '\u2026': '...',  # …
            '\u00A0': ' ',   # non-breaking space
            '\u200B': '',    # zero-width space
            '\u200C': '',    # zero-width non-joiner
            '\u200D': '',    # zero-width joiner
            '\uFEFF': '',    # BOM
            '\u2003': ' ',   # em space
            '\u2002': ' ',   # en space
            '\u2009': ' ',   # thin space
            # Arrows & symbols
            '\u2192': '->',  # →
            '\u2190': '<-',  # ←
            '\u2194': '<->', # ↔
            '\u2713': 'v',   # ✓
            '\u2714': 'v',   # ✔
            '\u2715': 'x',   # ✕
            '\u2716': 'x',   # ✖
            '\u2717': 'x',   # ✗
            '\u2718': 'x',   # ✘
            '\u00B7': '-',   # ·
            '\u25CF': '-',   # ●
            '\u25CB': 'o',   # ○
            '\u25AA': '-',   # ▪
            '\u25AB': '-',   # ▫
            '\u2605': '*',   # ★
            '\u2606': '*',   # ☆
        }
        for old, new in replacements.items():
            text = text.replace(old, new)
        # Final fallback: encode to latin-1 and drop anything that still can't be encoded
        try:
            text.encode('latin-1')
        except UnicodeEncodeError:
            text = text.encode('latin-1', errors='replace').decode('latin-1')
        return text

    def convert_docx_to_pdf(self, docx_buffer: io.BytesIO) -> io.BytesIO:
        """
        Convert DOCX buffer to PDF using FPDF2 (pure Python, no system dependencies).
        Reads DOCX content and recreates it as PDF with proper formatting.
        """
        if not FPDF2_AVAILABLE:
            raise ImportError(
                "fpdf2 is not installed. Install it with: pip install fpdf2==2.7.*"
            )
        
        from docx import Document as DocxDocument
        
        # Read DOCX from buffer
        docx_buffer.seek(0)
        doc = DocxDocument(docx_buffer)
        
        # Create PDF using FPDF2
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=19.05)  # 0.75 inch = 19.05mm margins
        # Calculate usable width: A4 width (210mm) minus left and right margins (2 * 19.05mm)
        page_width_mm = 210 - (2 * 19.05)  # Usable width after margins = 171.9mm
        
        # FPDF2 uses built-in fonts - Helvetica is closest to Arial
        # Font sizes in FPDF2 are in points (correct)
        # Spacing and cell heights need to be in mm
        
        # Convert points to mm for spacing (1 pt = 0.352778 mm)
        def pt_to_mm(pt_value):
            return pt_value * 0.352778
        
        # Process paragraphs from DOCX
        list_counter = 1  # For numbered lists (Learning Objectives)
        success_criteria_counter = 1  # Separate counter for Success Criteria
        in_success_criteria = False  # Track if we're in Success Criteria section
        pdf.add_page()
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                # Add small spacing for empty paragraphs
                pdf.ln(pt_to_mm(6))  # 6pt spacing
                continue
            
            # Check paragraph properties
            is_centered = para.alignment == WD_ALIGN_PARAGRAPH.CENTER
            has_runs = para.runs and len(para.runs) > 0
            is_bold = False
            font_size = None
            
            if has_runs:
                try:
                    first_run = para.runs[0]
                    is_bold = getattr(first_run, 'bold', False)
                    if hasattr(first_run, 'font') and first_run.font:
                        font_size = getattr(first_run.font, 'size', None)
                except (AttributeError, IndexError):
                    pass
            
            # Check if it's a title (centered, bold, 24pt or 20pt)
            if is_centered and has_runs and is_bold:
                try:
                    if font_size is not None and font_size in (self.TITLE_SIZE, Pt(24)):
                        actual_pt = 24 if font_size == Pt(24) else 20
                        pdf.ln(pt_to_mm(6))  # Space before title (6pt)
                        pdf.set_font('Helvetica', 'B', actual_pt)
                        cell_height = pt_to_mm(actual_pt * 1.2)
                        clean_text = self._sanitize_for_pdf(text)
                        pdf.cell(0, cell_height, clean_text, ln=1, align='C')
                        pdf.ln(pt_to_mm(4))  # Space after title
                        list_counter = 1
                        continue
                except Exception:
                    pass
            
            # Check if it's a header (bold, 14pt)
            if has_runs and is_bold:
                try:
                    if font_size is not None and font_size == self.HEADER_SIZE:
                        pdf.ln(pt_to_mm(6))  # Space before section heading (6pt)
                        pdf.set_font('Helvetica', 'B', 14)
                        cell_height = pt_to_mm(17)  # ~6mm for 14pt font
                        clean_text = self._sanitize_for_pdf(text)
                        pdf.cell(0, cell_height, clean_text, ln=1)
                        pdf.ln(pt_to_mm(4))  # Space after section heading (4pt)
                        # Reset counters based on section
                        if 'Success Criteria' in clean_text:
                            in_success_criteria = True
                            success_criteria_counter = 1
                        else:
                            in_success_criteria = False
                            list_counter = 1  # Reset list counter for Learning Objectives
                        continue
                except Exception:
                    pass
            
            # Handle lists
            style_name = getattr(para.style, 'name', '') if para.style else ''
            if style_name == 'List Bullet':
                pdf.set_font('Helvetica', '', 12)
                cell_height = pt_to_mm(14)
                pdf.cell(5, cell_height, '-', ln=0)
                clean_text = self._sanitize_for_pdf(text)
                pdf.multi_cell(page_width_mm - 5, cell_height, clean_text, ln=1)
                pdf.ln(pt_to_mm(2))
                continue  # Prevent fallthrough to other handlers
            
            # Check for manually numbered items (like "1. Text" for Success Criteria)
            import re
            text_stripped = text.strip()
            if text_stripped and re.match(r'^\d+\.\s+', text_stripped):
                pdf.set_font('Helvetica', '', 12)
                cell_height = pt_to_mm(14)
                match = re.match(r'^(\d+)\.\s*(.+)$', text_stripped)
                if match:
                    number = match.group(1)
                    content_text = match.group(2)
                    number_width = pdf.get_string_width(f'{number}.') + 2
                    pdf.cell(number_width, cell_height, f'{number}.', ln=0)
                    clean_text = self._sanitize_for_pdf(content_text)
                    pdf.multi_cell(page_width_mm - number_width, cell_height, clean_text, ln=1)
                    pdf.ln(pt_to_mm(2))
                continue
            elif style_name == 'List Number':
                cell_height = pt_to_mm(14)
                number_text = f'{list_counter}.'
                pdf.set_font('Helvetica', '', 12)
                number_width = pdf.get_string_width(number_text) + 2
                indent_mm = 0.5 * 25.4  # 0.5 in → mm
                pdf.cell(indent_mm - number_width, cell_height, '', ln=0)
                pdf.cell(number_width, cell_height, number_text, ln=0)
                clean_text = self._sanitize_for_pdf(text)
                pdf.multi_cell(page_width_mm - indent_mm, cell_height, clean_text, ln=1)
                list_counter += 1
            else:
                # Regular paragraph – render each run individually to preserve per-run bold
                cell_height = pt_to_mm(14)
                if has_runs and len(para.runs) > 1:
                    remaining_width = page_width_mm
                    for run in para.runs:
                        run_bold = getattr(run, 'bold', False)
                        style_flag = 'B' if run_bold else ''
                        pdf.set_font('Helvetica', style_flag, 12)
                        run_text = self._sanitize_for_pdf(run.text or '')
                        tw = pdf.get_string_width(run_text)
                        if tw <= remaining_width:
                            pdf.cell(tw, cell_height, run_text, ln=0)
                            remaining_width -= tw
                        else:
                            pdf.multi_cell(remaining_width, cell_height, run_text, ln=1)
                            remaining_width = page_width_mm
                    pdf.ln(cell_height)
                else:
                    pdf.set_font('Helvetica', '', 12)
                    clean_text = self._sanitize_for_pdf(text)
                    pdf.multi_cell(page_width_mm, cell_height, clean_text, ln=1)
                    pdf.ln(pt_to_mm(4))
        
        # Generate PDF to buffer
        pdf_buffer = io.BytesIO()
        pdf.output(pdf_buffer)
        pdf_buffer.seek(0)
        return pdf_buffer
