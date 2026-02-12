"""
DOCX Export for Lesson Starter
Exports to .docx while preserving template styles.
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_to_docx(
    lesson_text: str,
    output_path: str,
    template_path: str = None
) -> str:
    """
    Export lesson starter to DOCX file with template styles.
    
    Args:
        lesson_text: Generated lesson starter text
        output_path: Path to save DOCX file
        template_path: Optional path to .docx template
    
    Returns:
        Path to saved file
    """
    
    # Load template or create new document
    if template_path:
        doc = Document(template_path)
        # Clear existing body content but preserve styles
        for element in doc.element.body:
            doc.element.body.remove(element)
    else:
        doc = Document()
    
    # Parse lesson text into sections
    lines = lesson_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Add blank line as empty paragraph with List Number style
            p = doc.add_paragraph()
            try:
                p.style = 'List Number'
            except KeyError:
                pass  # Style doesn't exist, use default
            continue
        
        # Title: "Lesson Starter"
        if line == 'Lesson Starter':
            p = doc.add_paragraph(line)
            try:
                p.style = 'Title'
            except KeyError:
                # Fallback: make it look like a title
                run = p.runs[0]
                run.font.size = Pt(26)
                run.font.bold = True
            continue
        
        # Headers: Grade Level, Topic, Time Needed, section headers
        if line.startswith('Grade Level:') or line.startswith('Topic:') or line.startswith('Time Needed:'):
            p = doc.add_paragraph()
            try:
                p.style = 'List Number'
            except KeyError:
                pass
            
            # Split into label and value
            parts = line.split(':', 1)
            if len(parts) == 2:
                # Label bold, value normal
                run1 = p.add_run(parts[0] + ':')
                run1.bold = True
                run2 = p.add_run(parts[1])
                run2.bold = False
            else:
                p.add_run(line)
            continue
        
        # Section headers (no colons)
        if line in ['Key Lesson Ideas to Explore', 'Prior Knowledge to Connect', 'Teacher Opening Script']:
            p = doc.add_paragraph(line)
            try:
                p.style = 'List Number'
            except KeyError:
                pass
            run = p.runs[0]
            run.bold = True
            continue
        
        # Final question line (contains **...?**)
        if re.match(r'\*\*.*\?\*\*\$', line):
            # Strip asterisks but keep bold
            question_text = re.sub(r'\*\*', '', line)
            p = doc.add_paragraph()
            try:
                p.style = 'List Number'
            except KeyError:
                pass
            run = p.add_run(question_text)
            run.bold = True
            continue
        
        # Regular body text
        p = doc.add_paragraph(line)
        try:
            p.style = 'List Number'
        except KeyError:
            pass
    
    # Save document
    doc.save(output_path)
    return output_path


def export_lesson_starter_to_docx(
    lesson_text: str,
    output_path: str,
    template_path: str = None,
    title_style: str = "Title",
    body_style: str = "Normal"
) -> str:
    """
    Export lesson starter to DOCX with enhanced formatting options.
    
    Args:
        lesson_text: Generated lesson starter text
        output_path: Path to save DOCX file
        template_path: Optional template path
        title_style: Style name for title
        body_style: Style name for body text
    
    Returns:
        Path to saved file
    """
    return export_to_docx(lesson_text, output_path, template_path)
