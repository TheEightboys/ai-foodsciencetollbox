"""
Discussion Questions DOCX Export
Exports discussion questions to properly formatted Word documents.
"""

from __future__ import annotations
from docx import Document
from docx.shared import Pt

TITLE_STYLE = "Title"
BODY_STYLE = "Normal"

def _clear_document_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        body.remove(child)

def export_discussion_questions_docx(
    *,
    generated_text: str,
    template_path: str,
    output_path: str,
) -> None:
    doc = Document(template_path)
    _clear_document_body(doc)

    lines = [ln.rstrip() for ln in generated_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")]

    for ln in lines:
        if ln.strip() == "":
            doc.add_paragraph("", style=BODY_STYLE)
            continue

        # Accept both "DISCUSSION QUESTIONS" and "Discussion Questions"
        if ln.strip() in ("DISCUSSION QUESTIONS", "Discussion Questions"):
            doc.add_paragraph(ln.strip(), style=TITLE_STYLE)
            continue

        if ln.startswith("Grade Level:"):
            p = doc.add_paragraph("", style=BODY_STYLE)
            r = p.add_run("Grade Level:")
            r.bold = True
            p.add_run(ln[len("Grade Level:"):])
            continue

        if ln.startswith("Topic:"):
            p = doc.add_paragraph("", style=BODY_STYLE)
            r = p.add_run("Topic:")
            r.bold = True
            p.add_run(ln[len("Topic:"):])
            continue

        if ln.startswith("Teacher cue:"):
            p = doc.add_paragraph("", style=BODY_STYLE)
            r = p.add_run("Teacher cue:")
            r.bold = True
            p.add_run(ln[len("Teacher cue:"):])
            continue

        doc.add_paragraph(ln, style=BODY_STYLE)

    doc.save(output_path)


def export_discussion_questions_to_docx(
    questions_text: str,
    output_path: str,
    template_path: str = None,
    title_style: str = "Title",
    body_style: str = "Normal"
) -> str:
    """
    Export discussion questions to DOCX with enhanced formatting options.
    
    Args:
        questions_text: Generated discussion questions text
        output_path: Path to save DOCX file
        template_path: Optional template path
        title_style: Style name for title
        body_style: Style name for body text
    
    Returns:
        Path to saved file
    """
    # Load template or create new document
    if template_path:
        doc = Document(template_path)
        # Clear existing body content but preserve styles
        _clear_document_body(doc)
    else:
        doc = Document()
    
    # Parse discussion questions text
    lines = questions_text.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Add blank line
            p = doc.add_paragraph()
            try:
                p.style = body_style
            except KeyError:
                pass
            continue
        
        # Title: "DISCUSSION QUESTIONS" or "Discussion Questions"
        if line in ['DISCUSSION QUESTIONS', 'Discussion Questions']:
            p = doc.add_paragraph(line)
            try:
                p.style = title_style
            except KeyError:
                # Fallback: make it look like a title
                run = p.runs[0]
                run.font.size = Pt(26)
                run.font.bold = True
            continue
        
        # Headers: Grade Level, Topic, Teacher cue
        if line.startswith('Grade Level:') or line.startswith('Topic:') or line.startswith('Teacher cue:'):
            p = doc.add_paragraph()
            try:
                p.style = body_style
            except KeyError:
                pass
            
            # Split into label and value
            parts = line.split(':', 1)
            if len(parts) == 2:
                # Label bold, value normal
                run1 = p.add_run(parts[0] + ':')
                run1.bold = True
                run2 = p.add_run(parts[1])
                run2.bold = False
            else:
                p.add_run(line)
            continue
        
        # Regular body text
        p = doc.add_paragraph(line)
        try:
            p.style = body_style
        except KeyError:
            pass
    
    # Save document
    doc.save(output_path)
    return output_path
