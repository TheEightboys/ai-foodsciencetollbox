"""Test the updated format_learning_objectives to verify it matches the template."""
import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'Backend'))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings.development')

# Minimal Django setup
import django
django.setup()

from apps.generators.document_formatter import DocumentFormatter

# Sample content that mimics what the AI would return
sample_content = """1. Define what an enzyme is and explain its role as a biological catalyst.
2. Describe how enzymes interact with substrates using the concepts of active site and enzyme-substrate complex.
3. Identify everyday examples of enzyme activity in foods such as browning, ripening, and fermentation.
4. Explain how factors such as temperature and pH affect enzyme activity.
5. Compare denatured enzymes with active enzymes and describe what causes denaturation.
6. Analyze a simple food-related scenario and predict how enzyme activity will change under different conditions."""

formatter = DocumentFormatter()
result = formatter.format_learning_objectives(
    topic='Enzymes',
    grade_level='9th grade',
    content=sample_content,
    subject='Food Science'
)

# Save DOCX to inspect
docx_buf = result['docx']
docx_buf.seek(0)
output_path = 'test_output_objectives.docx'
with open(output_path, 'wb') as f:
    f.write(docx_buf.read())
print(f"DOCX saved to {output_path}")

# Now inspect the generated DOCX
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import lxml.etree as ET

doc = Document(output_path)
print(f"\nTotal paragraphs: {len(doc.paragraphs)}")

for i, p in enumerate(doc.paragraphs):
    text = p.text[:80] if p.text else '(empty)'
    style = p.style.name if p.style else 'None'
    align = p.alignment
    fmt = p.paragraph_format
    print(f"\nPara {i}: style={style}, align={align}")
    print(f"  Text: {repr(text)}")
    print(f"  space_before={fmt.space_before}, space_after={fmt.space_after}")
    print(f"  left_indent={fmt.left_indent}, first_line={fmt.first_line_indent}")
    
    for j, run in enumerate(p.runs):
        f = run.font
        print(f"  Run {j}: bold={f.bold}, name={f.name}, size={f.size}, text={repr(run.text[:50])}")

# Also test PDF conversion
docx_buf.seek(0)
try:
    pdf_buf = formatter.convert_docx_to_pdf(docx_buf)
    pdf_path = 'test_output_objectives.pdf'
    with open(pdf_path, 'wb') as f:
        f.write(pdf_buf.read())
    print(f"\nPDF saved to {pdf_path}")
except Exception as e:
    print(f"\nPDF conversion error: {e}")
